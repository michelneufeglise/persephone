"""
Intelligent Document Processing engine for Persephone.

Supports: PDF, DOCX, XLSX, CSV, TXT, MD, PNG/JPEG.
Pipeline: upload → extract text + images → store → on-demand operations
  (OCR, Q&A, summarize, classify, translate, extract entities, tables, export).
Operations that need vision call the user-configured vision/OCR model in Ollama.
"""

from __future__ import annotations

import base64
import io
import json
import logging
import mimetypes
import os
import re
import shutil
import tempfile
import time
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable

import httpx

log = logging.getLogger("idp_engine")

OLLAMA_BASE = os.getenv("OLLAMA_HOST", "http://localhost:11434")

# Model names whose prefix indicates vision/image-capable inference
VISION_PREFIXES = (
    "qwen2.5vl", "qwen2-vl", "qwen-vl", "qwen3vl",
    "minicpm-v", "minicpm-o", "openbmb/minicpm",
    "llama3.2-vision", "llama4-vision",
    "llava",
    "gemma3",
    "granite3-vision", "granite3.2-vision",
    "moondream", "bakllava", "cogvlm", "internvl",
)


def _name_is_vision(model: str) -> bool:
    lower = model.lower()
    # match either start-of-string or after a slash (e.g. "openbmb/minicpm-v:8b")
    base = lower.split(":", 1)[0]
    return any(base.startswith(p) or f"/{p}" in base for p in VISION_PREFIXES)


async def _list_installed_models() -> list[str]:
    try:
        async with httpx.AsyncClient(timeout=4.0) as c:
            r = await c.get(f"{OLLAMA_BASE}/api/tags")
            return [m["name"] for m in r.json().get("models", [])]
    except Exception:
        return []


async def _ollama_has_model(model: str) -> bool:
    installed = await _list_installed_models()
    if model in installed:
        return True
    # Allow tag-less matches (user picked "qwen2.5vl:7b", Ollama has "qwen2.5vl:7b-q4")
    base = model.split(":", 1)[0]
    return any(m.split(":", 1)[0] == base for m in installed)


_SIZE_RE = re.compile(r"(\d+\.?\d*)b\b", re.IGNORECASE)


def _estimate_params_b(name: str) -> float:
    """Pull the parameter-count (B) out of a model tag, default to 999B if unknown."""
    m = _SIZE_RE.search(name)
    if m:
        try:
            return float(m.group(1))
        except ValueError:
            pass
    return 999.0


async def find_installed_vision_model() -> str:
    """
    Return the smallest installed vision-capable model.
    Smaller wins so the fallback fits in memory alongside the user's chat model.
    """
    visions = [m for m in await _list_installed_models() if _name_is_vision(m)]
    if not visions:
        return ""
    visions.sort(key=_estimate_params_b)
    return visions[0]

# Storage root — writable per Electron app data dir in production builds.
from paths import uploads_dir
STORAGE_DIR = uploads_dir()


# ── Document model & registry ──────────────────────────────────────────────────
@dataclass
class Document:
    id:         str
    filename:   str
    mime:       str
    size:       int
    uploaded_at: float
    pages:      int = 0
    text:       str = ""
    page_texts: list[str] = field(default_factory=list)
    page_images: list[str] = field(default_factory=list)   # absolute paths to per-page PNGs
    meta:       dict[str, Any] = field(default_factory=dict)

    def to_dict(self, include_text: bool = False) -> dict:
        d = {
            "id":         self.id,
            "filename":   self.filename,
            "mime":       self.mime,
            "size":       self.size,
            "uploaded_at": int(self.uploaded_at * 1000),
            "pages":      self.pages,
            "preview":    self.text[:240] if self.text else "",
            "has_images": bool(self.page_images),
            "meta":       self.meta,
        }
        if include_text:
            d["text"]       = self.text
            d["page_texts"] = self.page_texts
        return d


# In-memory document registry, persisted to disk as JSON for crash safety
REGISTRY: dict[str, Document] = {}
REGISTRY_FILE = STORAGE_DIR / "_registry.json"


def _save_registry() -> None:
    serializable = {
        did: {
            **doc.to_dict(include_text=True),
            "page_images": doc.page_images,
            "uploaded_at": doc.uploaded_at,
        }
        for did, doc in REGISTRY.items()
    }
    REGISTRY_FILE.write_text(json.dumps(serializable))


def _load_registry() -> None:
    if not REGISTRY_FILE.exists():
        return
    try:
        data = json.loads(REGISTRY_FILE.read_text())
        for did, d in data.items():
            REGISTRY[did] = Document(
                id=d["id"], filename=d["filename"], mime=d["mime"], size=d["size"],
                uploaded_at=d["uploaded_at"] if isinstance(d["uploaded_at"], float) else d["uploaded_at"] / 1000,
                pages=d.get("pages", 0),
                text=d.get("text", ""),
                page_texts=d.get("page_texts", []),
                page_images=d.get("page_images", []),
                meta=d.get("meta", {}),
            )
    except Exception as exc:
        log.warning("Failed to load registry: %s", exc)


_load_registry()


def list_documents() -> list[dict]:
    return sorted(
        [d.to_dict() for d in REGISTRY.values()],
        key=lambda x: x["uploaded_at"], reverse=True,
    )


def get_document(doc_id: str) -> Document | None:
    return REGISTRY.get(doc_id)


def delete_document(doc_id: str) -> bool:
    doc = REGISTRY.pop(doc_id, None)
    if not doc:
        return False
    doc_dir = STORAGE_DIR / doc_id
    if doc_dir.exists():
        shutil.rmtree(doc_dir, ignore_errors=True)
    _save_registry()
    return True


# ── File extraction ────────────────────────────────────────────────────────────
def _detect_mime(filename: str) -> str:
    mime, _ = mimetypes.guess_type(filename)
    if not mime:
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        mime_map = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "csv": "text/csv",
            "md": "text/markdown",
            "txt": "text/plain",
            "png": "image/png",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
        }
        mime = mime_map.get(ext, "application/octet-stream")
    return mime


def _extract_pdf(path: Path, doc_dir: Path) -> tuple[list[str], list[str]]:
    """Extract per-page text and per-page PNG renders from a PDF."""
    import fitz   # type: ignore[import-not-found]

    page_texts: list[str] = []
    page_images: list[str] = []
    pdf = fitz.open(path)
    for i, page in enumerate(pdf):
        page_texts.append(page.get_text("text"))
        png_path = doc_dir / f"page_{i + 1:04d}.png"
        pix = page.get_pixmap(dpi=120)
        pix.save(png_path)
        page_images.append(str(png_path))
    pdf.close()
    return page_texts, page_images


def _extract_docx(path: Path) -> list[str]:
    from docx import Document as DocxDocument   # type: ignore[import-not-found]
    doc = DocxDocument(path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            paras.append("\t".join(cell.text.strip() for cell in row.cells))
    return ["\n".join(paras)] if paras else [""]


def _extract_xlsx(path: Path) -> tuple[list[str], dict]:
    from openpyxl import load_workbook  # type: ignore[import-not-found]
    wb = load_workbook(path, data_only=True, read_only=True)
    pages: list[str] = []
    sheets_meta = {}
    for ws_name in wb.sheetnames:
        ws = wb[ws_name]
        rows: list[str] = []
        row_count = 0
        col_count = 0
        for row in ws.iter_rows(values_only=True):
            row_count += 1
            if row:
                col_count = max(col_count, len(row))
                rows.append("\t".join("" if c is None else str(c) for c in row))
        sheets_meta[ws_name] = {"rows": row_count, "cols": col_count}
        pages.append(f"### Sheet: {ws_name}\n\n" + "\n".join(rows))
    wb.close()
    return pages, {"sheets": sheets_meta}


def _extract_csv(path: Path) -> list[str]:
    import pandas as pd  # type: ignore[import-not-found]
    try:
        df = pd.read_csv(path)
        return [df.to_string(index=False, max_rows=10000)]
    except Exception:
        return [path.read_text(errors="ignore")]


def _extract_image(path: Path, doc_dir: Path) -> list[str]:
    """Single-page 'document' for raw images. OCR happens on demand."""
    out = doc_dir / "page_0001.png"
    shutil.copy(path, out)
    return [str(out)]


async def ingest_file(filename: str, data: bytes) -> Document:
    """Persist a file, extract its text/images, register it."""
    doc_id = uuid.uuid4().hex[:12]
    doc_dir = STORAGE_DIR / doc_id
    doc_dir.mkdir(parents=True, exist_ok=True)

    raw_path = doc_dir / filename
    raw_path.write_bytes(data)

    mime = _detect_mime(filename)
    page_texts: list[str] = []
    page_images: list[str] = []
    meta: dict[str, Any] = {}

    try:
        if mime == "application/pdf":
            page_texts, page_images = _extract_pdf(raw_path, doc_dir)
        elif mime.endswith("wordprocessingml.document") or filename.lower().endswith(".docx"):
            page_texts = _extract_docx(raw_path)
        elif mime.endswith("spreadsheetml.sheet") or filename.lower().endswith(".xlsx"):
            page_texts, sheet_meta = _extract_xlsx(raw_path)
            meta.update(sheet_meta)
        elif mime == "text/csv" or filename.lower().endswith(".csv"):
            page_texts = _extract_csv(raw_path)
        elif mime.startswith("image/"):
            page_images = _extract_image(raw_path, doc_dir)
            page_texts  = [""]    # image-only — OCR fills this in on demand
        elif mime.startswith("text/"):
            page_texts = [raw_path.read_text(errors="ignore")]
        else:
            page_texts = [raw_path.read_text(errors="ignore")]
    except Exception as exc:
        log.exception("ingest failed: %s", exc)
        page_texts = [f"[ingest error: {exc}]"]

    full_text = "\n\n".join(page_texts).strip()

    doc = Document(
        id=doc_id, filename=filename, mime=mime, size=len(data),
        uploaded_at=time.time(),
        pages=max(len(page_texts), len(page_images)),
        text=full_text,
        page_texts=page_texts,
        page_images=page_images,
        meta=meta,
    )
    REGISTRY[doc_id] = doc
    _save_registry()
    return doc


# ── Vision-LLM helpers ────────────────────────────────────────────────────────
async def _ollama_vision_call(
    model: str, prompt: str, image_paths: list[str],
    *, num_predict: int = 2048,
) -> str:
    """Call Ollama with images. Resolves & validates the model before sending."""
    # Resolve the actual model to use ─────────────────────────────────────
    chosen = (model or "").strip()
    if not chosen or not _name_is_vision(chosen):
        fallback = await find_installed_vision_model()
        if fallback:
            log.info("Vision call: '%s' is text-only — using installed vision model '%s'",
                     chosen or "(none)", fallback)
            chosen = fallback
        else:
            raise RuntimeError(
                "No vision-capable model is installed. "
                "Install one with `ollama pull qwen2.5vl:7b` (or `minicpm-v`) "
                "and select it in Settings → Models."
            )

    if not await _ollama_has_model(chosen):
        raise RuntimeError(
            f"Model '{chosen}' is configured but not installed in Ollama. "
            f"Run `ollama pull {chosen}` from the terminal, then try again."
        )

    # Encode images to base64 ─────────────────────────────────────────────
    images_b64: list[str] = []
    for p in image_paths[:16]:        # cap to avoid blowing out context
        try:
            data = Path(p).read_bytes()
            if data:
                images_b64.append(base64.b64encode(data).decode())
        except Exception as exc:
            log.warning("Skipped unreadable image %s: %s", p, exc)

    if not images_b64:
        raise RuntimeError("No readable images to process")

    payload = {
        "model":  chosen,
        "prompt": prompt,
        "images": images_b64,
        "stream": False,
        "options": {
            "temperature": 0.2,
            "num_predict": num_predict,
            "num_thread":  10,
        },
    }

    try:
        async with httpx.AsyncClient(timeout=300.0) as client:
            r = await client.post(f"{OLLAMA_BASE}/api/generate", json=payload)
            if r.status_code >= 400:
                body_text = (r.text or "")[:300]
                log.error("Ollama %d on vision call (model=%s): %s",
                          r.status_code, chosen, body_text)
                raise RuntimeError(
                    f"Vision model '{chosen}' failed (Ollama HTTP {r.status_code}). "
                    f"{body_text or 'Check that the model accepts images.'}"
                )
            return (r.json().get("response") or "").strip()
    except httpx.RequestError as exc:
        raise RuntimeError(f"Could not reach Ollama at {OLLAMA_BASE}: {exc}")


async def _ollama_text_call(
    model: str, prompt: str, *, num_predict: int = 2048,
) -> str:
    if not model:
        # Fall back to whatever chat model is installed
        installed = await _list_installed_models()
        chat_candidates = [m for m in installed
                           if not any(p in m.lower() for p in ("embed", "orpheus"))]
        if not chat_candidates:
            raise RuntimeError(
                "No chat model is configured or installed. "
                "Install one with `ollama pull qwen2.5:7b` first."
            )
        model = chat_candidates[0]
        log.info("Text call: no model configured — using '%s'", model)

    if not await _ollama_has_model(model):
        raise RuntimeError(
            f"Model '{model}' is not installed in Ollama. "
            f"Run `ollama pull {model}` from the terminal, then try again."
        )

    try:
        async with httpx.AsyncClient(timeout=300.0) as client:
            r = await client.post(
                f"{OLLAMA_BASE}/api/generate",
                json={
                    "model": model, "prompt": prompt, "stream": False,
                    "options": {"temperature": 0.3, "num_predict": num_predict, "num_thread": 10},
                },
            )
            if r.status_code >= 400:
                body_text = (r.text or "")[:300]
                log.error("Ollama %d on text call (model=%s): %s",
                          r.status_code, model, body_text)
                raise RuntimeError(
                    f"Model '{model}' failed (Ollama HTTP {r.status_code}). {body_text}"
                )
            return (r.json().get("response") or "").strip()
    except httpx.RequestError as exc:
        raise RuntimeError(f"Could not reach Ollama at {OLLAMA_BASE}: {exc}")


# ── IDP operations ────────────────────────────────────────────────────────────
async def run_ocr(doc: Document, model: str, page_range: tuple[int, int] | None = None) -> str:
    """Run OCR on all (or a slice of) the document's page images."""
    if not doc.page_images:
        return doc.text or "(no images available for OCR)"

    images = doc.page_images
    if page_range:
        s, e = page_range
        images = images[max(0, s - 1) : e]

    prompt = (
        "Perform OCR on the provided image(s). Extract every readable text element exactly as written, "
        "preserving line breaks and reading order. Output ONLY the extracted text — no commentary."
    )
    text = await _ollama_vision_call(model, prompt, images, num_predict=4096)
    # cache result
    doc.meta["last_ocr_at"] = int(time.time() * 1000)
    doc.meta["ocr_model"]   = model
    if not doc.text:
        doc.text = text
    _save_registry()
    return text


async def summarize(doc: Document, model: str, style: str = "brief") -> str:
    style_prompts = {
        "brief":    "in 3-5 sentences",
        "detailed": "in 2-3 paragraphs with key facts and conclusions",
        "bullets":  "as a bullet list of the most important points",
    }
    prompt = (
        f"Summarize the following document {style_prompts.get(style, style_prompts['brief'])}.\n\n"
        f"--- DOCUMENT ---\n{doc.text[:32000]}\n--- END ---"
    )
    return await _ollama_text_call(model, prompt, num_predict=1024)


async def qa(doc: Document, model: str, question: str) -> str:
    prompt = (
        "Answer the user's question using ONLY the document below. "
        "If the answer is not in the document, say so plainly.\n\n"
        f"--- DOCUMENT ---\n{doc.text[:32000]}\n--- END ---\n\n"
        f"Question: {question}\nAnswer:"
    )
    return await _ollama_text_call(model, prompt, num_predict=1024)


async def extract_tables(doc: Document, model: str) -> list[dict]:
    """Return tables as a list of {title, rows: [[...]]} objects."""
    use_vision = bool(doc.page_images) and not (doc.text or "").strip()
    if use_vision:
        prompt = (
            "Extract all tables from the image(s). For each table, output strict JSON with this schema:\n"
            "{\"tables\": [{\"title\": \"...\", \"headers\": [...], \"rows\": [[...]]}]}\n"
            "If no tables are present, output {\"tables\": []}. Output ONLY JSON."
        )
        raw = await _ollama_vision_call(model, prompt, doc.page_images, num_predict=4096)
    else:
        prompt = (
            "Extract all tables from the following document into strict JSON:\n"
            "{\"tables\": [{\"title\": \"...\", \"headers\": [...], \"rows\": [[...]]}]}\n"
            f"Output ONLY JSON.\n\n--- DOCUMENT ---\n{doc.text[:32000]}\n--- END ---"
        )
        raw = await _ollama_text_call(model, prompt, num_predict=2048)

    # Extract the first {...} block
    m = re.search(r"\{[\s\S]*\}", raw)
    if not m:
        return []
    try:
        return json.loads(m.group(0)).get("tables", [])
    except Exception:
        return []


async def extract_entities(doc: Document, model: str) -> dict:
    """Pull dates, people, organizations, amounts, emails, URLs, addresses."""
    prompt = (
        "Extract all named entities from this document. Return strict JSON with these keys:\n"
        "{\"people\":[], \"organizations\":[], \"dates\":[], \"amounts\":[], "
        "\"emails\":[], \"urls\":[], \"addresses\":[], \"phone_numbers\":[]}\n"
        "Each list contains de-duplicated strings exactly as they appear. Output ONLY JSON.\n\n"
        f"--- DOCUMENT ---\n{doc.text[:32000]}\n--- END ---"
    )
    raw = await _ollama_text_call(model, prompt, num_predict=2048)
    m = re.search(r"\{[\s\S]*\}", raw)
    try:
        return json.loads(m.group(0)) if m else {}
    except Exception:
        return {}


async def classify(doc: Document, model: str) -> dict:
    prompt = (
        "Classify this document. Return strict JSON: "
        "{\"type\":\"invoice|contract|resume|report|email|letter|receipt|form|other\", "
        "\"language\":\"en|fr|...\", \"confidence\":0.0-1.0, \"topics\":[]}\n"
        "Output ONLY JSON.\n\n"
        f"--- DOCUMENT ---\n{doc.text[:16000]}\n--- END ---"
    )
    raw = await _ollama_text_call(model, prompt, num_predict=512)
    m = re.search(r"\{[\s\S]*\}", raw)
    try:
        return json.loads(m.group(0)) if m else {}
    except Exception:
        return {}


async def translate(doc: Document, model: str, target_lang: str) -> str:
    prompt = (
        f"Translate the following document into {target_lang}. Preserve formatting and structure. "
        "Output ONLY the translation.\n\n"
        f"--- DOCUMENT ---\n{doc.text[:24000]}\n--- END ---"
    )
    return await _ollama_text_call(model, prompt, num_predict=4096)


async def redact(doc: Document, model: str, categories: list[str]) -> str:
    cats = ", ".join(categories) if categories else "personal names, addresses, phone numbers, emails, IDs"
    prompt = (
        f"Redact the following from the document by replacing each occurrence with [REDACTED]: {cats}.\n"
        "Output the full document with redactions applied.\n\n"
        f"--- DOCUMENT ---\n{doc.text[:24000]}\n--- END ---"
    )
    return await _ollama_text_call(model, prompt, num_predict=4096)


# ── Exports ────────────────────────────────────────────────────────────────────
def export_markdown(doc: Document) -> bytes:
    md = f"# {doc.filename}\n\n"
    for i, page in enumerate(doc.page_texts, 1):
        if doc.pages > 1:
            md += f"\n## Page {i}\n\n"
        md += page + "\n"
    return md.encode()


def export_txt(doc: Document) -> bytes:
    return doc.text.encode() if doc.text else b""


def export_pdf(doc: Document) -> bytes:
    from reportlab.lib.pagesizes import letter  # type: ignore[import-not-found]
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak, Spacer
    from reportlab.lib.units import inch

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=0.7*inch, rightMargin=0.7*inch,
                            topMargin=0.7*inch, bottomMargin=0.7*inch)
    styles = getSampleStyleSheet()
    story: list[Any] = [Paragraph(f"<b>{doc.filename}</b>", styles["Title"]), Spacer(1, 12)]
    for i, page in enumerate(doc.page_texts, 1):
        if doc.pages > 1:
            story.append(Paragraph(f"<b>Page {i}</b>", styles["Heading2"]))
        for paragraph in page.split("\n\n"):
            para_html = paragraph.replace("\n", "<br/>").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if para_html.strip():
                story.append(Paragraph(para_html, styles["BodyText"]))
                story.append(Spacer(1, 6))
        if i < doc.pages:
            story.append(PageBreak())
    pdf.build(story)
    return buf.getvalue()


def export_xlsx(tables: list[dict]) -> bytes:
    from openpyxl import Workbook  # type: ignore[import-not-found]
    wb = Workbook()
    wb.remove(wb.active)
    for i, t in enumerate(tables or [], 1):
        title = (t.get("title") or f"Table {i}")[:31] or f"Table {i}"
        ws = wb.create_sheet(title=title)
        headers = t.get("headers", [])
        if headers:
            ws.append(headers)
        for row in t.get("rows", []):
            ws.append(row)
    if not wb.sheetnames:
        wb.create_sheet("Empty")
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def export_json(payload: Any) -> bytes:
    return json.dumps(payload, indent=2, ensure_ascii=False).encode()


def export_csv(tables: list[dict]) -> bytes:
    """Export the first table to CSV."""
    if not tables:
        return b""
    t = tables[0]
    headers = t.get("headers", [])
    rows = t.get("rows", [])
    lines: list[str] = []
    if headers:
        lines.append(",".join(_csv_quote(str(h)) for h in headers))
    for row in rows:
        lines.append(",".join(_csv_quote(str(c)) for c in row))
    return ("\n".join(lines) + "\n").encode()


def _csv_quote(s: str) -> str:
    if any(c in s for c in ',"\n'):
        return '"' + s.replace('"', '""') + '"'
    return s
